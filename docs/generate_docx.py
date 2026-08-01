#!/usr/bin/env python3
"""Generates docs/PROJECT_DOCUMENTATION.docx from README.md.

Kept as a script (not a one-off) so the Word doc can be regenerated whenever
README.md changes, instead of the two silently drifting apart. Run:

    python docs/generate_docx.py

Requires `python-docx` (not in requirements.txt — this is a documentation
tool, not a runtime dependency of the inspection system itself).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
README_PATH = PROJECT_ROOT / "README.md"
OUTPUT_PATH = PROJECT_ROOT / "docs" / "PROJECT_DOCUMENTATION.docx"

HEADING_COLOR = RGBColor(0x21, 0x02, 0x35)


def add_title_page(doc: Document) -> None:
    title = doc.add_heading("Fastener Defect & Assembly-Safety Inspector", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Project Documentation\n"
        "Computer-vision fastener inspection with an auditable, rule-based "
        "safety verdict"
    )
    run.font.size = Pt(14)
    run.font.italic = True

    doc.add_page_break()


def add_image_if_exists(doc: Document, rel_path: str, width_in: float = 6.0) -> bool:
    img_path = PROJECT_ROOT / rel_path
    if not img_path.exists():
        return False
    doc.add_picture(str(img_path), width=Inches(width_in))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def parse_markdown_to_docx(doc: Document, markdown_text: str) -> None:
    """Minimal Markdown -> docx renderer covering exactly what README.md
    uses: headings (#-####), fenced code blocks, tables, bullet lists,
    inline images, bold/italic/inline-code, and horizontal rules. This is
    intentionally not a general Markdown parser — it exists to keep the
    Word doc and the README from drifting apart, not as a reusable library.
    """
    lines = markdown_text.split("\n")
    i = 0
    in_code_block = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Fenced code blocks
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                _add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Skip the table of contents block (redundant in a Word doc which
        # has its own navigation via Word's heading styles / nav pane).
        if line.strip() == "## Table of contents":
            i += 1
            while i < len(lines) and not lines[i].startswith("---"):
                i += 1
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.*)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = _strip_inline_markdown(heading_match.group(2))
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.color.rgb = HEADING_COLOR
            i += 1
            continue

        # Markdown tables
        if line.strip().startswith("|") and i + 1 < len(lines) and re.match(
            r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]
        ):
            table_lines = [line]
            i += 2  # skip the header + separator row
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        # Inline images: ![alt](path)
        image_match = re.match(r"^!\[[^\]]*\]\(([^)]+)\)\s*$", line.strip())
        if image_match:
            added = add_image_if_exists(doc, image_match.group(1))
            if not added:
                doc.add_paragraph(f"[image not found: {image_match.group(1)}]")
            i += 1
            continue

        # Bullet list items
        bullet_match = re.match(r"^(\s*)-\s+(.*)$", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = _strip_inline_markdown(bullet_match.group(2))
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            doc.add_paragraph(text, style=style)
            i += 1
            continue

        # Numbered list items
        numbered_match = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if numbered_match:
            text = _strip_inline_markdown(numbered_match.group(1))
            doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Plain paragraph
        doc.add_paragraph(_strip_inline_markdown(line))
        i += 1


def _strip_inline_markdown(text: str) -> str:
    """Strip bold/italic/inline-code/link markers, keeping the plain text.
    Word styling for bold/italic is intentionally not reconstructed run-by-run
    here — the emphasis is preserved in wording, not font weight, since this
    document mirrors README.md's *content*, not its exact typography.
    """
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"\1", text)
    text = re.sub(r"`([^`]+?)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def _add_code_block(doc: Document, code: str) -> None:
    if not code.strip():
        return
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)


def _add_table(doc: Document, table_lines: list[str]) -> None:
    def split_row(row: str) -> list[str]:
        cells = row.strip().strip("|").split("|")
        return [_strip_inline_markdown(c) for c in cells]

    header = split_row(table_lines[0])
    rows = [split_row(r) for r in table_lines[1:]]

    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, header):
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text


def main() -> None:
    if not README_PATH.exists():
        raise SystemExit(f"{README_PATH} not found")

    markdown_text = README_PATH.read_text()

    doc = Document()
    for style_name in ("Normal",):
        doc.styles[style_name].font.size = Pt(11)

    add_title_page(doc)
    parse_markdown_to_docx(doc, markdown_text)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
